import streamlit as st
from pypdf import PdfReader
from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import os

# --- 1. INITIALIZATION ---
if 'ai_response' not in st.session_state:
    st.session_state['ai_response'] = ""
if 'model_answer' not in st.session_state:
    st.session_state['model_answer'] = ""

# Pulls the key from Streamlit Cloud Secrets instead of your code
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
client = genai.Client(api_key=GEMINI_API_KEY)

st.set_page_config(page_title="9618 CS AI Extractor Pro", layout="wide")


# --- 2. HELPERS ---
def get_working_model():
    try:
        for m in client.models.list():
            if 'generateContent' in m.supported_actions and "flash" in m.name.lower():
                return m.name
        return "gemini-1.5-flash"
    except:
        return "gemini-1.5-flash"


def highlight_keywords(text, keywords_str):
    if not keywords_str: return text
    keywords = [k.strip() for k in keywords_str.split(",") if k.strip()]
    highlighted_text = text
    for word in keywords:
        pattern = re.compile(re.escape(word), re.IGNORECASE)
        highlighted_text = pattern.sub(
            f'<span style="background-color: #FFFF00; color: black; padding: 2px;">{word}</span>', highlighted_text)
    return highlighted_text


def create_docx(text, topic, title_prefix="Questions"):
    doc = Document()
    doc.add_heading(f"9618 CS {title_prefix}: {topic}", 0)
    for line in text.split('\n'):
        if line.strip(): doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- 3. SIDEBAR ---
with st.sidebar:
    st.header("Search Parameters")
    topic = st.text_input("Syllabus Topic", "Data Representation")
    keywords = st.text_input("Keywords", "Binary")
    st.divider()
    year = st.selectbox("Year", [2021, 2022, 2023, 2024, 2025])
    month = st.selectbox("Month", ["June (s)", "Nov (w)", "March (m)"])
    paper = st.selectbox("Paper", [1, 2, 3, 4])
    variant = st.text_input("Variant", "1")
    if st.button("🔍 SEARCH SERVER"):
        st.session_state.model_answer = ""  # Clear old answers
        m_code = {"June (s)": "s", "Nov (w)": "w", "March (m)": "m"}[month]
        filename = f"9618_{m_code}{str(year)[2:]}_qp_{paper}{variant}.pdf"
        filepath = os.path.join("past_papers", filename)

        if os.path.exists(filepath):
            with st.spinner("Extracting..."):
                try:
                    reader = PdfReader(filepath)
                    full_text = "".join([p.extract_text() for p in reader.pages if p.extract_text()])[:20000]
                    prompt = f"Extract all questions about {topic} from this 9618 paper: {full_text}"
                    response = client.models.generate_content(model=get_working_model(), contents=prompt)
                    st.session_state.ai_response = response.text
                except Exception as e:
                    st.error(f"Error: {e}")
        else:
            st.error(f"File '{filename}' not found.")

# --- 4. MAIN DISPLAY ---
st.title("📂 9618 CS Automated Paper Search")

if st.session_state.ai_response:
    st.markdown("### Extracted Questions")
    st.markdown(highlight_keywords(st.session_state.ai_response, keywords), unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("🟦 Save Questions (Word)", create_docx(st.session_state.ai_response, topic),
                           f"{topic}_Questions.docx")

    with col2:
        if st.button("🧠 GENERATE MODEL ANSWERS"):
            with st.spinner("Solving..."):
                ans_prompt = f"Provide a concise Cambridge-style mark scheme for these questions: {st.session_state.ai_response}"
                ans_res = client.models.generate_content(model=get_working_model(), contents=ans_prompt)
                st.session_state.model_answer = ans_res.text

    if st.session_state.model_answer:
        st.divider()
        st.subheader("Suggested Mark Scheme")
        st.write(st.session_state.model_answer)
        st.download_button("🟩 Save Answers (Word)", create_docx(st.session_state.model_answer, topic, "Answers"),

                           f"{topic}_Answers.docx")

# --- FOOTER ---
st.markdown("---")

# Using a single container with centered alignment
st.markdown(
    """
    <div style="text-align: center; width: 100%;">
        <p style="font-size: 20px; font-weight: bold; margin-bottom: 5px;">
            ✨ PTES 9618 Computer Science PYP Resource Portal ✨
        </p>
        <p style="font-size: 16px; font-weight: bold; letter-spacing: 0.5px;">
            <span style="color: #FF0000;">🔴 Academically Excellence</span> | 
            <span style="color: #FFD700;">🟡 Future Readiness</span> | 
            <span style="color: #0070FF;">🔵 Digital & Integrity</span> | 
            <span style="color: #28A745;">🟢 Holistic & Growth</span>
        </p>
        <p style="color: gray; font-size: 14px; margin-top: 10px;">
            Creator: Miss Hajah Nurul Haziqah HN (PTES CS Tutor)
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

# --- FOOTER & VISITOR COUNTER ---
